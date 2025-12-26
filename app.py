# To run this code you need to install the following dependencies:
# pip install google-genai python-dotenv flask flask-cors pdfplumber python-docx pandas Pillow pytesseract openpyxl

import json
import os
import io
import base64
import re
from datetime import datetime
from flask import Flask, request, Response, jsonify
from flask_cors import CORS
from dotenv import load_dotenv
from google import genai
from google.genai import types
import pdfplumber
import docx
import pandas as pd
from PIL import Image
try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import InputFormat
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False

# .env dosyasından değişkenleri yükle
load_dotenv()

app = Flask(__name__)
CORS(app)  # Frontend'den istekler için CORS desteği

# Maksimum dosya boyutu: 10MB
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024


def generate_recommendations(inputs):
    """Gemini API kullanarak ürün önerileri oluştur"""
    
    # Prompt oluştur
    prompt = f"""
Sen bir ziraat karar-destek asistanısın.
Amaç: Verilen toprak + iklim + kısıt parametrelerine göre en uygun ürünleri öner.

ÖNEMLİ: Tüm parametreler OPSİYONEL'dir. Eğer bir parametre null/None ise, o parametre verilmemiş demektir.

KULLANILAN PARAMETRELER (girdi alanları - hepsi opsiyonel):
A) Toprak: soil_texture, pH, ec, organic_matter, nitrogen_N, phosphorus_P, potassium_K, lime_caCO3, cec
B) İklim: avg_temp_c, min_temp_c, max_temp_c, rainfall_mm, humidity_pct, drought_index
C) Konum/Zaman: country, province, district, lat (enlem), lon (boylam), season (mevsim), month (ay)
   - lat ve lon verilmişse, bu koordinatlara göre bölgenin iklim özelliklerini dikkate al
   - season ve month verilmişse, ekim zamanlaması için kullan
   - Eğer sadece lat/lon verilmişse, o bölgenin tipik iklim verilerini varsay
D) Kısıtlar: irrigation, previous_crop, goal

KURALLAR:
- Cevap Türkçe olacak.
- Çıktı SADECE JSON olacak (başka açıklama yazma).
- JSON şeması:
  {{
    "primary_crop": "string",
    "alternatives": ["string", "string", "string"],
    "confidence": 0-100,
    "reasons": ["..."],
    "risks": ["..."],
    "quick_actions": ["..."],
    "missing_inputs": ["..."],
    "assumptions": ["..."]
  }}
- Eğer bazı girdiler null/None ise, bunları "missing_inputs" içine yaz.
- Eksik veriler için makul varsayımlar yap ve bunları "assumptions" içine yaz.
- Verilen parametrelere göre en uygun ürün önerisini yap.
- Önerilerde 'goal' ve 'irrigation' alanlarını (varsa) mutlaka dikkate al.
- 'previous_crop' verilmişse münavebe mantığıyla aynı ürün tekrarına temkinli yaklaş.
- Sadece verilen parametrelere göre değerlendirme yap, eksik olanlar için varsayım yap.
- Eğer lat (enlem) ve lon (boylam) verilmişse, bu koordinatlara göre:
  * Bölgenin iklim özelliklerini (sıcaklık, yağış, nem) dikkate al
  * Bölgenin rakım ve topoğrafya özelliklerini değerlendir
  * O bölgeye özgü tarım uygulamalarını öner
  * Mevsim (season) ve ay (month) bilgisi varsa, ekim zamanlaması için kullan

GİRDİ (JSON - null değerler verilmemiş parametreleri gösterir):
{json.dumps(inputs, ensure_ascii=False, indent=2)}
"""
    
    client = genai.Client(
        api_key=os.environ.get("GEMINI_API_KEY"),
    )

    model = "gemma-3-27b-it"
    contents = [
        types.Content(
            role="user",
            parts=[
                types.Part.from_text(text=prompt),
            ],
        ),
    ]
    # Gemma modeli system_instruction ve GoogleSearch tool'unu desteklemiyor
    # Basit config kullanıyoruz
    generate_content_config = types.GenerateContentConfig()

    # Streaming response için generator
    def generate():
        full_text = ""
        try:
            for chunk in client.models.generate_content_stream(
                model=model,
                contents=contents,
                config=generate_content_config,
            ):
                if chunk.text:
                    full_text += chunk.text
                    yield chunk.text
            
            # Son kontrol - eğer JSON parse edilebilirse döndür
            try:
                json.loads(full_text)
            except:
                pass
        except Exception as e:
            # Hata durumunda kullanıcıya anlamlı mesaj gönder
            error_msg = str(e)
            if "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg or "quota" in error_msg.lower():
                yield json.dumps({
                    "error": "API quota aşıldı",
                    "message": "Gemini API kullanım limitiniz dolmuş. Lütfen planınızı ve faturalama detaylarınızı kontrol edin.",
                    "details": "https://ai.google.dev/gemini-api/docs/rate-limits"
                }, ensure_ascii=False)
            elif "401" in error_msg or "UNAUTHENTICATED" in error_msg:
                yield json.dumps({
                    "error": "API key hatası",
                    "message": "Geçersiz veya eksik API key. Lütfen .env dosyanızda GEMINI_API_KEY değişkenini kontrol edin."
                }, ensure_ascii=False)
            else:
                yield json.dumps({
                    "error": "API hatası",
                    "message": f"Bir hata oluştu: {error_msg}"
                }, ensure_ascii=False)
    
    return generate()


@app.route('/api/recommend', methods=['POST'])
def recommend():
    """API endpoint - form verilerini alıp öneri döndürür"""
    try:
        inputs = request.get_json()
        
        if not inputs:
            return jsonify({"error": "Input verisi bulunamadı"}), 400
        
        # Streaming response döndür
        return Response(
            generate_recommendations(inputs),
            mimetype='text/plain',
            headers={
                'Cache-Control': 'no-cache',
                'X-Accel-Buffering': 'no'
            }
        )
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def extract_text_from_pdf(file_content):
    """PDF dosyasından metin çıkar - pdfplumber kütüphanesi kullanılıyor (CPU'da hızlı çalışır)"""
    try:
        pdf = pdfplumber.open(io.BytesIO(file_content))
        text = ""
        for page_num, page in enumerate(pdf.pages, 1):
            # Normal metin
            page_text = page.extract_text() or ""
            text += f"\n--- Sayfa {page_num} ---\n"
            text += page_text
            
            # Tabloları da çıkar (çok önemli!)
            tables = page.extract_tables()
            if tables:
                text += f"\n--- Sayfa {page_num} Tabloları ---\n"
                for table_num, table in enumerate(tables, 1):
                    text += f"\nTablo {table_num}:\n"
                    for row in table:
                        if row:
                            # Satırı temizle ve birleştir
                            clean_row = [str(cell).strip() if cell else "" for cell in row]
                            text += " | ".join(clean_row) + "\n"
                    text += "\n"
        
        pdf.close()
        return text.strip()
    except Exception as e:
        return f"PDF okuma hatası: {str(e)}"


def extract_text_from_word(file_content):
    """Word dosyasından metin çıkar - python-docx kütüphanesi kullanılıyor (CPU'da hızlı çalışır)"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        
        # Paragrafları oku
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Tabloları da oku (çok önemli!)
        for table_num, table in enumerate(doc.tables, 1):
            text += f"\n--- Tablo {table_num} ---\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text += " | ".join(row_text) + "\n"
            text += "\n"
        
        return text.strip()
    except Exception as e:
        return f"Word okuma hatası: {str(e)}"


def extract_data_from_csv(file_content):
    """CSV/Excel dosyasından veri çıkar - pandas kütüphanesi kullanılıyor (CPU'da hızlı çalışır)"""
    try:
        # Önce CSV olarak dene
        try:
            df = pd.read_csv(io.BytesIO(file_content), encoding='utf-8')
        except:
            try:
                df = pd.read_csv(io.BytesIO(file_content), encoding='latin-1')
            except:
                df = pd.read_csv(io.BytesIO(file_content), encoding='iso-8859-9')
        
        # Sütun isimlerini normalize et (küçük harf, boşlukları temizle)
        df.columns = df.columns.str.strip().str.lower()
        
        # DataFrame'i detaylı text formatına çevir
        # Önce sütun isimlerini ve veri tiplerini ekle
        text = "CSV/Excel Dosyası İçeriği:\n\n"
        text += f"Sütunlar: {', '.join(df.columns.tolist())}\n"
        text += f"Satır Sayısı: {len(df)}\n\n"
        
        # Her satırı detaylı göster
        for idx, row in df.iterrows():
            text += f"Satır {idx + 1}:\n"
            for col in df.columns:
                value = row[col]
                if pd.notna(value):
                    text += f"  {col}: {value}\n"
            text += "\n"
        
        # DataFrame'in string gösterimini de ekle
        text += "\nTablo Görünümü:\n"
        text += df.to_string()
        
        return text
    except Exception as e:
        # Excel dosyası olabilir
        try:
            df = pd.read_excel(io.BytesIO(file_content))
            df.columns = df.columns.str.strip().str.lower()
            
            text = "Excel Dosyası İçeriği:\n\n"
            text += f"Sütunlar: {', '.join(df.columns.tolist())}\n"
            text += f"Satır Sayısı: {len(df)}\n\n"
            
            for idx, row in df.iterrows():
                text += f"Satır {idx + 1}:\n"
                for col in df.columns:
                    value = row[col]
                    if pd.notna(value):
                        text += f"  {col}: {value}\n"
                text += "\n"
            
            text += "\nTablo Görünümü:\n"
            text += df.to_string()
            
            return text
        except Exception as e2:
            return f"CSV/Excel okuma hatası: {str(e)} / {str(e2)}"


def extract_text_from_image_docling(file_content, mime_type):
    """Docling kullanarak resimden metin çıkar ve JSON formatına dönüştür"""
    if not DOCLING_AVAILABLE:
        return None
    
    try:
        # Geçici dosya oluştur
        import tempfile
        import os
        
        # Dosya uzantısını mime_type'a göre belirle
        suffix = '.jpg'
        if 'png' in mime_type:
            suffix = '.png'
        elif 'gif' in mime_type:
            suffix = '.gif'
        elif 'bmp' in mime_type:
            suffix = '.bmp'
        elif 'webp' in mime_type:
            suffix = '.webp'
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(file_content)
            tmp_path = tmp_file.name
        
        try:
            # Docling converter oluştur
            converter = DocumentConverter()
            
            # Resmi işle
            result = converter.convert(tmp_path)
            
            # Metni çıkar - Docling'in text özelliğini kullan
            extracted_text = ""
            if hasattr(result.document, 'text'):
                extracted_text = result.document.text
            elif hasattr(result.document, 'export_to_dict'):
                # Dict formatına çevir
                doc_json = result.document.export_to_dict()
                # Metni çıkar (text, tables, vb.)
                if 'content' in doc_json:
                    for item in doc_json['content']:
                        if 'text' in item:
                            extracted_text += item['text'] + "\n"
                        elif 'table' in item:
                            # Tablo varsa
                            if 'rows' in item['table']:
                                for row in item['table']['rows']:
                                    if 'cells' in row:
                                        row_text = " | ".join([cell.get('text', '') for cell in row['cells']])
                                        extracted_text += row_text + "\n"
            else:
                # Fallback: string'e çevir
                extracted_text = str(result.document)
            
            return extracted_text.strip() if extracted_text else None
        finally:
            # Geçici dosyayı sil
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    except Exception as e:
        # Hata durumunda None döndür (fallback için)
        return None


def extract_text_from_image_docling_or_gemini(file_content, mime_type):
    """Resimden metin çıkar - Önce Docling, sonra Gemini Vision API (fallback)
    
    NOT: 
    - PDF: pdfplumber (CPU'da hızlı)
    - Word: python-docx (CPU'da hızlı)
    - CSV/Excel: pandas (CPU'da hızlı)
    - Resim: Docling (öncelikli), Gemini Vision API (fallback)
    """
    # Önce Docling ile dene
    if DOCLING_AVAILABLE:
        docling_result = extract_text_from_image_docling(file_content, mime_type)
        if docling_result and len(docling_result) > 50:
            return docling_result
    
    # Docling başarısız olursa Gemini Vision API'ye fallback
    try:
        client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
        # Gemini Vision API için uygun model kullan
        model = "gemini-1.5-flash"
        
        prompt = """Bu resim bir toprak analiz raporu içeriyor. Lütfen TÜM metinleri, sayıları, değerleri ve bilgileri çok dikkatli bir şekilde çıkar.

ÖNEMLİ: Hiçbir bilgiyi atlama! Tüm tabloları, değerleri, tarihleri, isimleri oku.

ÇIKARILMASI GEREKEN BİLGİLER:
1. Numune bilgileri: Numune No/Kodu, Numune Alım Tarihi, Analiz Tarihi
2. Konum bilgileri: İl, İlçe, Saha/Parsel bilgisi
3. Laboratuvar: Laboratuvar adı, onay/kaşe/imza bilgisi
4. Toprak analiz değerleri: pH, EC, Organik Madde (%), Azot (N), Fosfor (P), Potasyum (K), Kireç (CaCO3), CEC
5. Mikro elementler: Kalsiyum (Ca), Magnezyum (Mg), Kükürt (S), Demir (Fe), Çinko (Zn), Mangan (Mn), Bakır (Cu), Bor (B)
6. Diğer: Toplam Tuz, SAR, ESP, Organik Karbon, Toprak Nemi, Bulk Density
7. Toprak bünyesi: Kumlu/Tınlı/Killi
8. Değerlendirme seviyesi: Düşük/Orta/Yüksek
9. Gübreleme önerisi: N-P-K formatında veya metin olarak

TÜM metni, sayıları, birimleri, tarihleri eksiksiz çıkar. Tabloları, listeleri, her şeyi oku."""
        
        try:
            contents = [
                types.Content(
                    role="user",
                    parts=[
                        types.Part.from_text(text=prompt),
                        types.Part.from_bytes(
                            data=file_content,
                            mime_type=mime_type
                        )
                    ],
                ),
            ]
            
            response = client.models.generate_content(
                model=model,
                contents=contents,
            )
            
            return response.text
        except Exception as e1:
            # Gemini Vision API başarısız olursa gemma-3-27b-it dene
            error_str = str(e1).lower()
            if "404" in error_str or "not found" in error_str or "not enabled" in error_str:
                # Model bulunamadı, gemma-3-27b-it dene
                try:
                    model = "gemma-3-27b-it"
                    response = client.models.generate_content(
                        model=model,
                        contents=contents,
                    )
                    return response.text
                except Exception as e2:
                    return f"Resim OCR hatası (Gemini): {str(e2)}"
            else:
                return f"Resim OCR hatası: {str(e1)}"
    except Exception as e:
        return f"Resim OCR genel hatası: {str(e)}"


def parse_extracted_text(text):
    """Çıkarılan metni Gemini API ile parse edip form verilerine çevir - Gelişmiş versiyon
    
    Daha güçlü model (gemini-1.5-flash) kullanarak veri parsing yapar.
    Eğer gemini-1.5-flash kullanılamazsa gemma-3-27b-it'e fallback yapar.
    """
    try:
        client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))
        
        # Belge parsing için gemma-3-27b-it modeli kullan
        model = "gemma-3-27b-it"
        
        prompt = f"""
Sen bir toprak analiz raporu parser'ısın. Verilen metinden TÜM toprak analiz verilerini çok dikkatli, detaylı ve eksiksiz bir şekilde çıkar ve JSON formatında döndür.

KRİTİK GÖREV: Metindeki TÜM bilgileri bulmaya çalış. Hiçbir değeri atlama! Tablolarda, paragraflarda, dipnotlarda, her yerde ara!

ÇALIŞMA PRENSİBİ:
1. Metni satır satır, kelime kelime oku
2. Her sayısal değeri, her tarihi, her ismi bul
3. Alternatif isimleri kontrol et (örn: "Fosfor" = "P" = "P2O5")
4. Birimleri kaldır ama değeri koru
5. Binlik ayırıcıları kaldır (1,250.5 -> 1250.5)
6. Virgüllü sayıları noktaya çevir (6,5 -> 6.5)
7. Yüzdeli değerlerde hassasiyeti koru (%2.345 -> 2.345)

ÇIKARILMASI GEREKEN ALANLAR VE ALTERNATİF İSİMLERİ:

ZORUNLU ALANLAR:
1. sample_code: "Numune No", "Numune Kodu", "Numune Numarası", "Örnek No", "Sample No", "Sample Code", "Numune ID"
2. sample_date: "Numune Alım Tarihi", "Örnek Alım Tarihi", "Numune Alma Tarihi", "Sample Date", "Alım Tarihi" (YYYY-MM-DD formatına çevir)
3. analysis_date: "Analiz Tarihi", "Tahlil Tarihi", "Test Tarihi", "Analysis Date", "Rapor Tarihi" (YYYY-MM-DD formatına çevir)
4. province: "İl", "Province", "İl Adı", "Şehir", "City", "Location"
5. sample_depth: "Numune Derinliği", "Örnek Derinliği", "Depth", "Derinlik", "cm" ile biten sayılar (sadece sayıyı al)
6. laboratory_name: "Laboratuvar", "Lab", "Laboratuvar Adı", "Laboratory", "Tahlil Yeri", "Analiz Yeri"
7. pH: "pH", "Ph", "PH", "pH değeri", "pH degeri" (0-14 arası sayı)
8. organic_matter: "Organik Madde", "OM", "Organik Materyal", "%" işareti ile birlikte olabilir (sadece sayıyı al)
9. phosphorus_P: "Fosfor", "P", "P2O5", "Phosphorus", "Alınabilir P", "Available P" (sadece sayıyı al)
10. potassium_K: "Potasyum", "K", "K2O", "Potassium", "Değişebilir K", "Exchangeable K" (sadece sayıyı al)

GENELLİKLE BULUNAN ALANLAR:
11. ec: "EC", "Elektriksel İletkenlik", "Electrical Conductivity", "Tuzluluk", "Salinity", "dS/m" ile birlikte olabilir
12. lime_caCO3: "Kireç", "CaCO3", "Lime", "Kireç %", "Kireç içeriği", "%" işareti ile birlikte olabilir
13. soil_texture: "Toprak Bünyesi", "Tekstür", "Texture", "Bünye" - Değerler: "kumlu", "tınlı", "killi", "kum", "kil", "tin"
14. nitrogen_N: "Azot", "N", "Nitrogen", "Toplam N", "Total N", "Alınabilir N", "Available N"
15. evaluation_level: "Değerlendirme", "Seviye", "Level", "Rating" - Değerler: "düşük", "orta", "yüksek", "low", "medium", "high"
16. fertilization_recommendation: "Gübreleme", "Fertilization", "Öneri", "Recommendation", "N-P-K", "NPK" formatında olabilir

OPSİYONEL MİKRO ELEMENTLER:
17. district: "İlçe", "District", "County", "İlçe Adı"
18. calcium_Ca: "Kalsiyum", "Ca", "Calcium", "Ca++", "Ca+2"
19. magnesium_Mg: "Magnezyum", "Mg", "Magnesium", "Mg++", "Mg+2"
20. sulfur_S: "Kükürt", "S", "Sulfur", "Sülfür"
21. iron_Fe: "Demir", "Fe", "Iron", "Fe++", "Fe+2", "Fe+3"
22. zinc_Zn: "Çinko", "Zn", "Zinc", "Zn++"
23. manganese_Mn: "Mangan", "Mn", "Manganese", "Mn++", "Mn+2"
24. copper_Cu: "Bakır", "Cu", "Copper", "Cu++", "Cu+2"
25. boron_B: "Bor", "B", "Boron"
26. cec: "CEC", "Katyon Değişim Kapasitesi", "Cation Exchange Capacity"
27. total_salt: "Toplam Tuz", "Total Salt", "Tuz", "Salt"
28. sar: "SAR", "Sodium Adsorption Ratio"
29. esp: "ESP", "Exchangeable Sodium Percentage"
30. organic_carbon_C: "Organik Karbon", "Organic Carbon", "C", "Carbon"
31. soil_moisture: "Toprak Nemi", "Soil Moisture", "Nem", "Moisture", "%" ile birlikte olabilir
32. bulk_density: "Bulk Density", "Hacim Ağırlığı", "Yoğunluk", "Density", "g/cm3" veya benzeri birimlerle

TARİH FORMATLARI (TÜM FORMATLARI TANIMLA):
- DD.MM.YYYY -> YYYY-MM-DD
- DD/MM/YYYY -> YYYY-MM-DD
- DD-MM-YYYY -> YYYY-MM-DD
- YYYY.MM.DD -> YYYY-MM-DD
- YYYY/MM/DD -> YYYY-MM-DD
- "15.03.2024" -> "2024-03-15"
- "15/03/2024" -> "2024-03-15"
- "2024.03.15" -> "2024-03-15"

SAYI FORMATLARI:
- "6.5", "6,5", "6.50" -> 6.5
- "% 2.3", "2.3%", "2,3 %" -> 2.3 (yüzde işaretini kaldır)
- "45 mg/kg", "45 mg kg-1", "45 mg/kg toprak" -> 45 (birimleri kaldır)
- "120 mm" -> 120 (birimleri kaldır)
- "1250.5", "1,250.5", "1250,5" -> 1250.5 (binlik ayırıcıları kaldır)
- "0.05", "0,05" -> 0.05 (ondalık değerler)
- "2.345%" -> 2.345 (yüzde işaretini kaldır, sayıyı koru - hassasiyeti koru)

ÖNEMLİ DEĞER ARALIKLARI (GERÇEKÇİ):
- pH: 0-14 arası (genellikle 4-9 arası)
- Organik Madde: 0-100% (genellikle 0.5-10% arası, virgüllü değerler olabilir: 2.345%)
- Fosfor (P): 0-5000 mg/kg (genellikle 5-200 mg/kg arası, ama 1000+ değerler de olabilir: 1250.5)
- Potasyum (K): 0-5000 mg/kg (genellikle 50-500 mg/kg arası, ama 1000+ değerler de olabilir: 2500)
- EC: 0-50 dS/m (genellikle 0-10 dS/m arası, ama daha yüksek değerler de olabilir: 15.5)
- Mikro elementler (Ca, Mg, Fe, Zn, Mn, Cu): 0-10000 mg/kg (genellikle 10-5000 mg/kg arası, 1000+ değerler normal)
- Bor (B): 0-1000 mg/kg veya ppm (genellikle 0.1-10 arası, virgüllü: 0.05)
- CEC: 0-500 meq/100g (genellikle 5-50 arası, ama 100+ değerler de olabilir: 150.5)

ÖNEMLİ KURALLAR:
1. Metindeki TÜM değerleri bul, eksik bırakma
2. Alternatif isimleri kontrol et (örn: "Fosfor" ve "P" aynı şey)
3. Birimleri kaldır (%, mg/kg, cm, mm, dS/m vb.)
4. Tarihleri mutlaka YYYY-MM-DD formatına çevir
5. Sayısal değerleri sayı olarak döndür (string değil)
6. Toprak bünyesi için: "kum" -> "kumlu", "kil" -> "killi", "tin" -> "tınlı"
7. Çıktı SADECE JSON olacak, başka açıklama yazma
8. Bulunamayan değerler için null kullan
9. Eğer bir değer farklı formatlarda yazılmışsa, hepsini kontrol et

ÖRNEK ÇIKTI FORMATI:
{{
  "sample_code": "NUM-2024-001",
  "sample_date": "2024-03-15",
  "analysis_date": "2024-03-20",
  "province": "Konya",
  "district": "Meram",
  "sample_depth": 30,
  "laboratory_name": "Toprak Analiz Laboratuvarı",
  "pH": 6.7,
  "organic_matter": 2.3,
  "phosphorus_P": 45,
  "potassium_K": 35,
  "ec": 1.2,
  "lime_caCO3": 5.5,
  "soil_texture": "tınlı",
  "nitrogen_N": 50,
  "evaluation_level": "orta",
  "fertilization_recommendation": "20-10-10",
  "calcium_Ca": 120,
  "magnesium_Mg": 45,
  "sulfur_S": 15,
  "iron_Fe": 8.5,
  "zinc_Zn": 2.3,
  "manganese_Mn": 12.5,
  "copper_Cu": 1.8,
  "boron_B": 0.5,
  "cec": 25.5,
  "total_salt": 0.8,
  "sar": 2.5,
  "esp": 5.2,
  "organic_carbon_C": 1.2,
  "soil_moisture": 15.5,
  "bulk_density": 1.35
}}

ÖNEMLİ NOTLAR VE ÖRNEKLER:
- Metinde "pH: 6.7" veya "pH = 6.7" veya "pH 6.7" görürsen -> "pH": 6.7
- Metinde "Organik Madde: %2.3" veya "OM: 2.3%" görürsen -> "organic_matter": 2.3 (yüzde işaretini kaldır)
- Metinde "Fosfor (P): 45 mg/kg" veya "P: 1250.5 mg/kg" görürsen -> "phosphorus_P": 45 veya 1250.5 (birimi kaldır, sayıyı koru)
- Metinde "Potasyum: 350 mg/kg" veya "K: 2500" görürsen -> "potassium_K": 350 veya 2500
- Metinde "EC: 1.25 dS/m" veya "Elektriksel İletkenlik: 2.5" görürsen -> "ec": 1.25 veya 2.5
- Metinde "Kalsiyum: 1250 mg/kg" veya "Ca: 2500" görürsen -> "calcium_Ca": 1250 veya 2500 (1000+ değerler olabilir)
- Metinde "CEC: 25.5 meq/100g" veya "Katyon Değişim: 150" görürsen -> "cec": 25.5 veya 150 (100+ değerler olabilir)
- Metinde "15.03.2024" veya "15/03/2024" görürsen -> "sample_date": "2024-03-15"
- Metinde "Konya İli" veya "Konya" görürsen -> "province": "Konya"
- Metinde "Tınlı" veya "Loam" görürsen -> "soil_texture": "tınlı"
- Metinde "Orta seviye" veya "Medium" görürsen -> "evaluation_level": "orta"
- Eğer bir değer farklı formatlarda yazılmışsa (örn: "pH 6.7" veya "pH=6.7" veya "pH:6.7"), hepsini tanı
- Tablolardaki değerleri de oku - özellikle sayısal değerler tablolarda olabilir
- Alt satırlarda, dipnotlarda, küçük yazılarda da bilgi olabilir, onları da oku
- Binlik ayırıcıları kaldır: "1,250.5" -> 1250.5
- Virgüllü sayıları noktaya çevir: "6,5" -> 6.5
- Yüzdeli değerlerde hassasiyeti koru: "%2.345" -> 2.345 (0.01 step için)

ÇIKARILACAK METİN:
{text}

JSON ÇIKTISI (SADECE JSON, başka hiçbir şey yazma. Tüm bulduğun değerleri ekle):
"""
        
        # Belge parsing için daha güçlü model kullan (gemini-1.5-flash)
        # Eğer kullanılamazsa gemma-3-27b-it'e fallback
        model = "gemini-1.5-flash"
        contents = [
            types.Content(
                role="user",
                parts=[
                    types.Part.from_text(text=prompt),
                ],
            ),
        ]
        
        generate_content_config = types.GenerateContentConfig(
            temperature=0.1,  # Daha deterministik sonuçlar için düşük temperature
        )
        
        try:
            response = client.models.generate_content(
                model=model,
                contents=contents,
                config=generate_content_config,
            )
        except Exception as e:
            # Eğer gemini-1.5-flash başarısız olursa (404, quota, erişim vb.), gemma-3-27b-it'e fallback
            error_str = str(e).lower()
            if "404" in error_str or "not found" in error_str or "quota" in error_str or "429" in error_str or "not enabled" in error_str or "not available" in error_str or "permission" in error_str:
                # Fallback to gemma model
                model = "gemma-3-27b-it"
                response = client.models.generate_content(
                    model=model,
                    contents=contents,
                    config=generate_content_config,
                )
            else:
                raise e
        
        # JSON parse et
        result_text = response.text.strip()
        # Markdown code block'ları temizle
        result_text = result_text.replace('```json', '').replace('```', '').strip()
        
        # JSON içindeki kısmı bul
        json_start = result_text.find('{')
        json_end = result_text.rfind('}') + 1
        if json_start != -1 and json_end > json_start:
            result_text = result_text[json_start:json_end]
        
        try:
            parsed_data = json.loads(result_text)
            
            # Veri normalizasyonu ve temizleme
            normalized_data = normalize_parsed_data(parsed_data)
            
            return normalized_data
        except json.JSONDecodeError as e:
            # JSON parse edilemezse, tekrar deneme yap
            # Bazen model JSON dışında açıklama da ekliyor
            try:
                # Sadece JSON kısmını bul
                import re
                json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
                if json_match:
                    parsed_data = json.loads(json_match.group())
                    normalized_data = normalize_parsed_data(parsed_data)
                    return normalized_data
            except:
                pass
            return {"error": "Metin parse edilemedi", "raw_text": result_text[:500]}
            
    except Exception as e:
        return {"error": f"Parse hatası: {str(e)}"}


def normalize_parsed_data(data):
    """Parse edilmiş verileri normalize et - tarih formatları, sayı formatları, birimler"""
    normalized = {}
    
    for key, value in data.items():
        if value is None or value == '':
            normalized[key] = None
            continue
        
        # Tarih alanları
        if 'date' in key.lower() and isinstance(value, str):
            normalized[key] = normalize_date(value)
        # Sayısal alanlar
        elif key in ['pH', 'organic_matter', 'ec', 'lime_caCO3', 'sample_depth', 
                     'phosphorus_P', 'potassium_K', 'nitrogen_N', 'calcium_Ca', 
                     'magnesium_Mg', 'sulfur_S', 'iron_Fe', 'zinc_Zn', 'manganese_Mn',
                     'copper_Cu', 'boron_B', 'cec', 'total_salt', 'sar', 'esp',
                     'organic_carbon_C', 'soil_moisture', 'bulk_density']:
            normalized[key] = normalize_number(value)
        # Toprak bünyesi normalizasyonu
        elif key == 'soil_texture' and isinstance(value, str):
            normalized[key] = normalize_soil_texture(value)
        # Değerlendirme seviyesi normalizasyonu
        elif key == 'evaluation_level' and isinstance(value, str):
            normalized[key] = normalize_evaluation_level(value)
        # Diğer string alanlar
        else:
            normalized[key] = str(value).strip() if value else None
    
    return normalized


def normalize_date(date_str):
    """Tarih string'ini YYYY-MM-DD formatına çevir"""
    if not date_str or not isinstance(date_str, str):
        return None
    
    date_str = date_str.strip()
    
    # Farklı tarih formatlarını dene
    formats = [
        '%Y-%m-%d', '%d.%m.%Y', '%d/%m/%Y', '%d-%m-%Y',
        '%Y.%m.%d', '%Y/%m/%d', '%Y-%m-%d',
        '%d.%m.%y', '%d/%m/%y', '%d-%m-%y',
        '%Y.%m.%d', '%Y/%m/%d'
    ]
    
    for fmt in formats:
        try:
            dt = datetime.strptime(date_str, fmt)
            return dt.strftime('%Y-%m-%d')
        except:
            continue
    
    # Regex ile tarih bulma
    patterns = [
        r'(\d{4})[.\-/](\d{1,2})[.\-/](\d{1,2})',  # YYYY-MM-DD
        r'(\d{1,2})[.\-/](\d{1,2})[.\-/](\d{4})',  # DD-MM-YYYY
        r'(\d{1,2})[.\-/](\d{1,2})[.\-/](\d{2})',  # DD-MM-YY
    ]
    
    for pattern in patterns:
        match = re.search(pattern, date_str)
        if match:
            groups = match.groups()
            if len(groups[0]) == 4:  # YYYY-MM-DD
                year, month, day = groups
            else:  # DD-MM-YYYY
                day, month, year = groups
                if len(year) == 2:
                    year = '20' + year
            
            try:
                return f"{year}-{int(month):02d}-{int(day):02d}"
            except:
                continue
    
    return None


def normalize_number(value):
    """Sayısal değeri normalize et - birimleri kaldır, virgülü noktaya çevir"""
    if value is None:
        return None
    
    if isinstance(value, (int, float)):
        return float(value)
    
    if not isinstance(value, str):
        try:
            return float(value)
        except:
            return None
    
    # String'den sayı çıkar
    value = value.strip()
    
    # Yüzde işaretini kaldır
    value = value.replace('%', '').strip()
    
    # Birimleri kaldır (mg/kg, dS/m, cm, mm, g/cm3 vb.)
    value = re.sub(r'\s*(mg/kg|mg kg-1|dS/m|cm|mm|g/cm3|kg/ha|ppm|meq/100g)\s*', '', value, flags=re.IGNORECASE)
    
    # Virgülü noktaya çevir
    value = value.replace(',', '.')
    
    # Sadece sayı ve nokta karakterlerini al
    value = re.sub(r'[^\d.]', '', value)
    
    try:
        return float(value) if value else None
    except:
        return None


def normalize_soil_texture(value):
    """Toprak bünyesi değerini normalize et"""
    if not value or not isinstance(value, str):
        return None
    
    value = value.lower().strip()
    
    # Eşleştirmeler
    mappings = {
        'kumlu': 'kumlu',
        'kum': 'kumlu',
        'sandy': 'kumlu',
        'sand': 'kumlu',
        'tınlı': 'tınlı',
        'tin': 'tınlı',
        'tinli': 'tınlı',
        'loam': 'tınlı',
        'loamy': 'tınlı',
        'killi': 'killi',
        'kil': 'killi',
        'clay': 'killi',
        'clayey': 'killi',
    }
    
    return mappings.get(value, value)


def normalize_evaluation_level(value):
    """Değerlendirme seviyesini normalize et"""
    if not value or not isinstance(value, str):
        return None
    
    value = value.lower().strip()
    
    mappings = {
        'düşük': 'düşük',
        'dusuk': 'düşük',
        'low': 'düşük',
        'orta': 'orta',
        'medium': 'orta',
        'yüksek': 'yüksek',
        'yuksek': 'yüksek',
        'high': 'yüksek',
    }
    
    return mappings.get(value, value)


@app.route('/api/upload-file', methods=['POST'])
def upload_file():
    """Dosya yükleme ve veri çıkarma endpoint'i"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "Dosya bulunamadı"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "Dosya seçilmedi"}), 400
        
        # Dosya içeriğini oku
        file_content = file.read()
        filename = file.filename.lower()
        mime_type = file.content_type
        
        # Dosya tipine göre işle
        extracted_text = ""
        extraction_method = ""
        
        try:
            if filename.endswith('.pdf'):
                extracted_text = extract_text_from_pdf(file_content)
                extraction_method = "PDF (pdfplumber - CPU)"
            elif filename.endswith(('.doc', '.docx')):
                extracted_text = extract_text_from_word(file_content)
                extraction_method = "Word (python-docx - CPU)"
            elif filename.endswith(('.csv', '.xlsx', '.xls')):
                extracted_text = extract_data_from_csv(file_content)
                extraction_method = "CSV/Excel (pandas - CPU)"
            elif filename.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')):
                # Önce Docling, sonra Gemini Vision API fallback
                extracted_text = extract_text_from_image_docling_or_gemini(file_content, mime_type)
                if DOCLING_AVAILABLE and not (extracted_text.startswith("Resim OCR") or "hatası" in extracted_text.lower()):
                    extraction_method = "Resim (Docling - CPU)"
                elif extracted_text.startswith("Resim OCR") or "hatası" in extracted_text.lower():
                    extraction_method = "Resim (Gemini Vision API - Bulut - Fallback)"
                else:
                    extraction_method = "Resim (Docling - CPU)"
            else:
                return jsonify({
                    "success": False,
                    "error": "Desteklenmeyen dosya formatı",
                    "message": "PDF, Word, CSV, Excel veya resim dosyası yükleyin."
                }), 400
        except Exception as e:
            return jsonify({
                "success": False,
                "error": "Dosya okuma hatası",
                "message": f"Dosya okunurken hata oluştu: {str(e)}"
            }), 400
        
        if not extracted_text or extracted_text.startswith("hata") or extracted_text.startswith("PDF okuma") or extracted_text.startswith("Word okuma") or extracted_text.startswith("CSV okuma") or extracted_text.startswith("Resim OCR"):
            return jsonify({
                "success": False,
                "error": "Dosya okunamadı",
                "message": extracted_text or "Dosya içeriği çıkarılamadı",
                "extraction_method": extraction_method
            }), 400
        
        # Çıkarılan metni parse et
        try:
            parsed_data = parse_extracted_text(extracted_text)
        except Exception as e:
            return jsonify({
                "success": False,
                "error": "Veri parse hatası",
                "message": f"Çıkarılan metin parse edilemedi: {str(e)}",
                "extracted_text_preview": extracted_text[:500]
            }), 400
        
        if "error" in parsed_data:
            return jsonify({
                "success": False,
                "error": "Veri parse edilemedi",
                "message": parsed_data.get("error", "Bilinmeyen hata"),
                "extracted_text_preview": extracted_text[:500],
                "extraction_method": extraction_method
            }), 400
        
        # Eşleşen alanları say
        matched_fields = [k for k, v in parsed_data.items() if v is not None and v != '']
        
        return jsonify({
            "success": True,
            "data": parsed_data,
            "extracted_text_preview": extracted_text[:200],
            "extraction_method": extraction_method,
            "matched_fields_count": len(matched_fields),
            "matched_fields": matched_fields
        })
    
    except Exception as e:
        return jsonify({"error": f"Sunucu hatası: {str(e)}"}), 500


@app.route('/')
def index():
    """Ana sayfa"""
    return "Tarım Asistanı API - /api/recommend ve /api/upload-file endpoint'lerini kullanın"


if __name__ == "__main__":
    app.run(debug=True, port=5001)

